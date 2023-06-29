
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

# Cria um novo arquivo PDF
c = canvas.Canvas("C:\\Users\\Everaldo Junior\\Documents\\Algorítmos - Autoria Web\\aulas-autoria-web\\PDF Python\\meu_documento.pdf", pagesize=letter)

# Adiciona texto
c.drawString(100, 700, "Olá, mundo!")

# Adiciona uma imagem
c.drawInlineImage("C:\\Users\\Everaldo Junior\\Documents\\Algorítmos - Autoria Web\\aulas-autoria-web\\PDF Python\\eu.jpeg", 100, 500)

# Finaliza o arquivo PDF
c.save()


"""
from docx import Document

document = Document()
document.add_heading('Tabela', level=1)

table = document.add_table(rows=3, cols=3)
table.cell(0, 0).text = "testeH"
table.cell(0, 1).text = "testeHHH"
table.cell(0, 2).text = "teste"
table.cell(1, 0).text = "teste"
table.cell(1, 1).text = "teste"
table.cell(1, 2).text = "teste"
table.cell(2, 0).text = "teste"
table.cell(2, 1).text = "teste"
table.cell(2, 2).text = "teste"

document.save("C:\\Users\\Everaldo Junior\\Documents\\Algorítmos - Autoria Web\\aulas-autoria-web\\PDF Python\\Autoria.pdf") #onde o documento vai ser salvo

#DOCUMENTAÇÃO PYTHON-DOCX: https://python-docx.readthedocs.io/en/latest/
"""