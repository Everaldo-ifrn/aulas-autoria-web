from docx import Document
from docx.shared import Cm

doc = Document() #Criando um objeto para o documento

doc.add_picture('C://Users//Everaldo Junior//Documents//Algorítmos - Autoria Web//aulas-autoria-web//PDF Python//ImagemNapne1.png', width=Cm(1.92), height=Cm(2.07)) #Adcionando 1ª fotocom largura, altura definidas
paragraph = doc.paragraphs[-1] #Selecionando e armazenando a imagem na variavel
paragraph.alignment = 1 #Aliando a imagem no centro [0 = esquerda; 1 = centro; 2 = direita]
doc.add_paragraph('Você é? __titulo__')

doc.save('C://Users//Everaldo Junior//Documents//Algorítmos - Autoria Web//aulas-autoria-web//PDF Python//DocxEstudoPython.docx') 